#!/usr/bin/env python3
"""Convert CRC-MMIS manuscript markdown to PeerJ DOCX with embedded figures"""

import re, os, io
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BASE = r"C:\Users\xingw\WorkBuddy\2026-06-04-22-17-17"
SRC = f"{BASE}/results/manuscript/CRC-MMIS-manuscript-FINAL.md"
DST = f"{BASE}/submission/CRC-MMIS-Manuscript-v3.5.docx"
FIGDIR = f"{BASE}/results/figures/final"
SUPPDIR = f"{BASE}/results/figures/cptac"
CELLCHATDIR = f"{BASE}/results/figures/cellchat"

TITLE = "Single-Cell Characterization of Malignant-Myeloid Interaction Transcriptomic Features Identifies a Dual-Circuit Adenosinergic Signaling Hypothesis in Colorectal Cancer"

with open(SRC, "r", encoding="utf-8") as f:
    text = f.read()

# Strip metadata header (everything before --- after title)
text = re.sub(r"^# .*?\n---+\n", "", text, flags=re.DOTALL)
text = re.sub(r"\*\*Target Journal:.*?\*\*Date:.*?\n", "", text)
text = re.sub(r"\n\*Manuscript prepared for.*?\*\*\n", "", text)
text = re.sub(r"\*\*Final QC\*\*.*?\n", "", text)
text = re.sub(r"\*\*Data Transparency Notes[\s\S]*?\*\*Final QC\*\*[\s\S]*?PeerJ requirements\.\n", "", text)
text = re.sub(r"> \*\*File:\*\* `.*?`\n", "", text)
text = re.sub(r"\n{3,}", "\n\n", text)

# Figure name mapping
fig_names = {
    "Figure 1": "Figure1_scRNA.jpg",
    "Figure 2": "Figure2_Validation.jpg",
    "Figure 3": "Figure3_CMS.jpg",
    "Figure 4": "Figure4_Adenosine.jpg",
    "Figure 5": "Figure5_TIDE.jpg",
    "Figure 6": "Figure6_MSI_Survival.jpg",
}

# Supplementary figure paths
supp_fig_paths = {
    "Figure S12": os.path.join(SUPPDIR, "FigureS12_CPTAC_Validation.jpg"),
    "Figure S10": os.path.join(CELLCHATDIR, "FigureS10_CellChat.jpg"),
    "Figure S11": os.path.join(CELLCHATDIR, "FigureS11_GSE205506_CellChat.jpg"),
}

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21.59)
section.page_height = Cm(27.94)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Line numbers
lnNumType = parse_xml(f'<w:lnNumType {nsdecls("w")} w:countBy="1"/>')
section._sectPr.append(lnNumType)

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run(TITLE)
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True
doc.add_paragraph()  # blank line

# Parse markdown
lines = text.split("\n")
i = 0
in_table = False
table_rows = []

def add_p(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = bold
    return p

def insert_figure(fig_num):
    """Insert figure image before the caption"""
    if fig_num in fig_names:
        path = os.path.join(FIGDIR, fig_names[fig_num])
        if os.path.exists(path):
            try:
                # Convert JPG to PNG via PIL for python-docx compatibility
                img = Image.open(path)
                png_bytes = io.BytesIO()
                img.save(png_bytes, format='PNG')
                png_bytes.seek(0)
                
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p_img.add_run()
                r.add_picture(png_bytes, width=Inches(5.5))
                doc.add_paragraph()
            except Exception as e:
                print(f"  Warning: Could not embed {fig_num}: {e}")

while i < len(lines):
    line = lines[i]
    
    # Figure heading → insert image first
    fig_match = re.match(r"^### (Figure \d)\. (.+)$", line)
    if fig_match:
        fig_num = fig_match.group(1)
        insert_figure(fig_num)
        add_p(f"{fig_num}. {fig_match.group(2)}", bold=True)
        i += 1
        # Skip file path line and caption content
        while i < len(lines):
            nl = lines[i]
            if nl.startswith("Caption:") or nl.startswith("**Caption:**"):
                caption_text = nl.replace("Caption:", "").replace("**Caption:**", "").replace("**", "").strip()
                # Read multi-line caption
                cap_lines = [caption_text]
                i += 1
                while i < len(lines) and lines[i].strip() and not lines[i].startswith("###") and not lines[i].startswith("---"):
                    if lines[i].strip():
                        cap_lines.append(lines[i].strip())
                    i += 1
                add_p(" ".join(cap_lines))
                break
            elif nl.startswith("###"):  # next section without caption
                break
            i += 1
        continue
    
    # Tables
    if line.startswith("| ") and "|" in line:
        cols = [c.strip() for c in line.split("|") if c.strip()]
        if not in_table:
            in_table = True
            table_rows = [cols]
        else:
            table_rows.append(cols)
        i += 1
        if i < len(lines) and re.match(r"^\|[-| ]+\|$", lines[i]):
            i += 1
        continue
    
    if in_table:
        if table_rows and len(table_rows) >= 2:
            ncols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=ncols)
            table.style = 'Table Grid'
            for ri, row_data in enumerate(table_rows):
                for ci, cell_text in enumerate(row_data):
                    if ci < ncols:
                        cell = table.rows[ri].cells[ci]
                        cell.text = cell_text
                        for pp in cell.paragraphs:
                            for rr in pp.runs:
                                rr.font.name = 'Times New Roman'
                                rr.font.size = Pt(9)
                                if ri == 0:
                                    rr.bold = True
        table_rows = []
        in_table = False
    
    # Separators
    if line.strip() == "---":
        i += 1
        continue
    
    # Empty line
    if not line.strip():
        i += 1
        continue
    
    # Headings
    stripped = line.strip()
    if stripped.startswith("### "):
        add_p(stripped[4:], bold=True)
    elif stripped.startswith("## "):
        add_p(stripped[3:], bold=True)
    elif stripped.startswith("**") and stripped.endswith("**"):
        add_p(stripped.strip("*"), bold=True)
    else:
        # Clean formatting marks
        cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
        cleaned = re.sub(r"<!--.*?-->", "", cleaned)
        if cleaned.strip():
            add_p(cleaned.strip())
    i += 1

doc.save(DST)
print(f"Saved with embedded figures: {DST}")
